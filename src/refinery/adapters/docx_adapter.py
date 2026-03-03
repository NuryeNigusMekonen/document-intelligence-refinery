from __future__ import annotations

from pathlib import Path

from ..models import ExtractedDocument, ExtractedPage, ProvenanceRef, TableObject, TextBlock


class DocxAdapter:
    name = "docx_adapter"

    def available(self) -> bool:
        try:
            import docx  # noqa: F401

            return True
        except Exception:
            return False

    def extract(self, file_path: Path, doc_id: str) -> ExtractedDocument:
        from docx import Document

        document = Document(str(file_path))
        section_stack: list[str] = []
        blocks: list[TextBlock] = []
        tables: list[TableObject] = []
        reading_order = 0

        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            paragraph_style = paragraph.style
            style_name = str(paragraph_style.name or "") if paragraph_style is not None else ""
            if style_name.startswith("Heading"):
                level = 1
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except Exception:
                    level = 1
                section_stack = section_stack[: max(level - 1, 0)] + [text]

            is_list = style_name.lower().startswith("list") or text[:2].strip().endswith(".")
            prefix = "[LIST] " if is_list else ""
            prov = ProvenanceRef(
                doc_name=file_path.name,
                ref_type="word_section",
                section_path=section_stack.copy() if section_stack else ["Document Root"],
                content_hash="pending",
            )
            blocks.append(
                TextBlock(
                    text=prefix + text,
                    bbox=(0.0, float(reading_order) * 10.0, 1000.0, float(reading_order + 1) * 10.0),
                    reading_order=reading_order,
                    confidence=0.9,
                    provenance=prov,
                )
            )
            reading_order += 1

        for t_idx, table in enumerate(document.tables):
            all_rows = [[(cell.text or "").strip() for cell in row.cells] for row in table.rows]
            if not all_rows:
                continue
            headers = all_rows[0]
            rows = all_rows[1:] if len(all_rows) > 1 else []
            prov = ProvenanceRef(
                doc_name=file_path.name,
                ref_type="word_section",
                section_path=section_stack.copy() if section_stack else ["Document Root"],
                content_hash="pending",
            )
            tables.append(
                TableObject(
                    bbox=(0.0, 0.0, 1000.0, 1000.0),
                    headers=headers,
                    rows=rows,
                    reading_order=10000 + t_idx,
                    confidence=0.88,
                    provenance=prov,
                )
            )

        page = ExtractedPage(page_number=1, width=1000.0, height=2000.0, blocks=blocks, tables=tables, figures=[])
        return ExtractedDocument(doc_id=doc_id, doc_name=file_path.name, pages=[page])
