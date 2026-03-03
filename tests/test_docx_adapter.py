from pathlib import Path

import pytest

from refinery.adapters.docx_adapter import DocxAdapter
from refinery.chunking import Chunker
from refinery.config import Settings
from refinery.storage import ArtifactStore


@pytest.mark.skipif(__import__("importlib").util.find_spec("docx") is None, reason="python-docx not installed")
def test_docx_adapter_and_section_provenance(tmp_path: Path):
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Financial Results", level=1)
    doc.add_paragraph("Revenue increased in FY2024.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Metric"
    tbl.rows[0].cells[1].text = "Value"
    tbl.rows[1].cells[0].text = "Revenue"
    tbl.rows[1].cells[1].text = "$10"
    doc.save(str(path))

    adapter = DocxAdapter()
    extracted = adapter.extract(path, "doc_docx")
    assert extracted.pages

    settings = Settings(workspace_root=tmp_path)
    store = ArtifactStore(settings)
    ldus = Chunker(settings, store).run(extracted)
    assert ldus
    assert any(ref.ref_type == "word_section" for ldu in ldus for ref in ldu.page_refs)
